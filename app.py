from flask import Flask, render_template, request, send_file, jsonify
import requests
from bs4 import BeautifulSoup
from datetime import datetime
import os
import sys
import webbrowser
from threading import Timer
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

app = Flask(__name__, template_folder=resource_path("templates"))

def open_browser():
    try:
        webbrowser.open_new("http://127.0.0.1:5001")
    except Exception:
        pass

def crawl_naver_news(query, max_items=50):
    results = []
    seen_titles = set()
    
    for start in range(1, max_items, 10):
        url = f"https://search.naver.com/search.naver?where=news&query={query}&start={start}"
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
        }
        
        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
        except Exception:
            continue

        soup = BeautifulSoup(response.text, 'html.parser')
        
        news_items = soup.select('a.news_tit')
        if not news_items:
            news_items = soup.select('.sds-comps-text-type-headline1')
        
        if not news_items:
            main_pack = soup.find(id="main_pack")
            if main_pack:
                news_items = main_pack.select('.sds-comps-text-type-headline1, .sds-comps-text-type-body2')
        
        for item in news_items:
            if len(results) >= max_items:
                break
                
            title = item.get_text(strip=True)
            if not title or title in seen_titles or len(title) < 5:
                continue
                
            link = item.get('href') if item.has_attr('href') else None
            if not link:
                parent_a = item.find_parent('a')
                if parent_a:
                    link = parent_a.get('href')
            
            snippet = ""
            parent_area = item.find_parent(class_='news_area')
            if not parent_area:
                parent_area = item.find_parent(class_='news_info')
            
            if parent_area:
                dsc = parent_area.select_one('.news_dsc, .api_txt_lines')
                if dsc:
                    snippet = dsc.get_text(strip=True)
            
            if not snippet:
                next_content = item.find_next(class_='sds-comps-text-type-body1')
                if next_content:
                    snippet = next_content.get_text(strip=True)

            source = "알 수 없음"
            info_group = item.find_parent(class_='news_info') or item.find_parent(class_='news_area')
            if info_group:
                info_tags = info_group.select('.info')
                for tag in info_tags:
                    text = tag.get_text(strip=True)
                    if text and '전' not in text and not any(c.isdigit() for c in text.split('.')):
                        source = text
                        break
                
                if source == "알 수 없음":
                    press_tag = info_group.select_one('.press')
                    if press_tag:
                        source = press_tag.get_text(strip=True)
            
            if source == "알 수 없음":
                source_tag = item.find_previous(class_='sds-comps-text-type-body2')
                if source_tag:
                    potential_source = source_tag.get_text(strip=True)
                    if '전' not in potential_source:
                        source = potential_source

            results.append({
                'title': title,
                'link': link if link else '#',
                'snippet': snippet if snippet else 'No summary available.',
                'source': source
            })
            seen_titles.add(title)
        
        if len(results) >= max_items:
            break
            
    return results

@app.route('/')
def index():
    return render_template('index.html')

def generate_head_message(results):
    if not results:
        return "수집된 기사가 없어 요약을 생성할 수 없습니다."
    
    top_results = [r for r in results[:5] if r['snippet'] != 'No summary available.']
    if not top_results:
        return f"현재 '{len(results)}'건의 뉴스가 검색되었습니다. 상세 내용은 아래 지면을 확인해 주시기 바랍니다."
    
    head_text = "【 오늘의 주요 뉴스 요약 】\n\n"
    list_items = []
    
    for r in top_results:
        snippet = r['snippet']
        first_sentence = snippet.split('. ')[0].strip()
        if not first_sentence.endswith('.'):
            first_sentence += '.'
            
        list_items.append(f'"{first_sentence}" - {r["source"]} -')
    
    return head_text + "\n".join(list_items)

@app.route('/crawl', methods=['POST'])
def crawl():
    query = request.form.get('query', 'ax ai')
    try:
        limit = int(request.form.get('limit', 50))
    except (ValueError, TypeError):
        limit = 50
    
    # Enforce limit between 10 and 100
    limit = max(10, min(limit, 100))
        
    results = crawl_naver_news(query, max_items=limit)
    head_message = generate_head_message(results)
    return jsonify({
        'results': results,
        'head_message': head_message
    })

@app.route('/download', methods=['POST'])
def download():
    data = request.json
    query = data.get('query', 'ax ai')
    results = data.get('results', [])
    
    if not results:
        return "No data", 400
        
    now = datetime.now()
    filename = f"{now.strftime('%Y-%m-%d-%H%M%S')}.txt"
    filepath = os.path.join(os.getcwd(), filename)
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"네이버 뉴스 검색 결과: '{query}'\n")
        f.write(f"수집 날짜: {now.strftime('%Y-%m-%d')}\n")
        f.write(f"수집 시간: {now.strftime('%H:%M:%S')}\n")
        f.write("-" * 50 + "\n\n")
        for item in results:
            f.write(f"Title: {item['title']}\n")
            f.write(f"Snippet: {item.get('snippet', '')}\n")
            f.write(f"Link: {item['link']}\n\n")
            
    return send_file(filepath, as_attachment=True)

@app.route('/download_word', methods=['POST'])
def download_word():
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    data = request.json
    query = data.get('query', 'ax ai')
    results = data.get('results', [])
    head_message = data.get('head_message', '')

    if not results:
        return "No data", 400

    doc = Document()
    
    title = doc.add_heading('THE NAVER DAILY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    now = datetime.now()
    date_para = doc.add_paragraph(now.strftime('%A, %B %d, %Y').upper())
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 50)

    doc.add_heading('Editorial Summary', level=1)
    summary_para = doc.add_paragraph(head_message)
    summary_para.style.font.italic = True
    doc.add_paragraph("-" * 50)

    for item in results:
        p = doc.add_paragraph()
        run = p.add_run(item['title'])
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph(item.get('snippet', ''))
        
        source_para = doc.add_paragraph(f"— {item.get('source', 'Naver News')} —")
        source_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        link_para = doc.add_paragraph()
        run_link = link_para.add_run(f"Link: {item['link']}")
        run_link.font.size = Pt(8)
        doc.add_paragraph()

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)

    filename = f"NAVER_DAILY_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    
    return send_file(
        target_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/send_email', methods=['POST'])
def send_email():
    data = request.json
    recipient = data.get('email')
    query = data.get('query')
    results = data.get('results', [])
    head_message = data.get('head_message', '')

    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port_raw = os.environ.get('SMTP_PORT', '587')
    smtp_user = os.environ.get('SMTP_USER')
    smtp_password = os.environ.get('SMTP_PASSWORD')

    if not all([smtp_server, smtp_user, smtp_password]):
        return jsonify({'error': 'SMTP server is not configured. Set SMTP_SERVER, SMTP_USER, and SMTP_PASSWORD environment variables.'}), 500
    
    server_host = str(smtp_server)
    server_port = int(smtp_port_raw)
    auth_user = str(smtp_user)
    auth_pw = str(smtp_password)

    try:
        now = datetime.now()
        subject = f"[THE NAVER DAILY] {query.upper()} EDITION - {now.strftime('%Y-%m-%d')}"
        
        html_body = f"""
        <div style="font-family: 'Georgia', serif; background-color: #f4f1ea; padding: 20px; color: #1a1a1a;">
            <div style="max-width: 600px; margin: 0 auto; border: 1px solid #1a1a1a; background-color: #ffffff; padding: 30px;">
                <h1 style="text-align: center; font-size: 36px; border-bottom: 4px double #1a1a1a; padding-bottom: 10px; margin-bottom: 10px;">THE NAVER DAILY</h1>
                <p style="text-align: center; font-weight: bold; margin-bottom: 20px; text-transform: uppercase;">SEOUL, KOREA | {now.strftime('%A, %B %d, %Y').upper()}</p>
                
                <div style="background-color: #efece4; padding: 15px; border: 1px solid #1a1a1a; font-style: italic; margin-bottom: 25px;">
                    <h3 style="text-align: center; margin-top: 0; text-transform: uppercase; font-size: 14px; border-bottom: 1px solid #1a1a1a;">Editorial Summary</h3>
                    <p style="white-space: pre-line;">{head_message}</p>
                </div>
                
                <div style="display: grid; grid-template-columns: 1fr; gap: 20px;">
        """
        
        for item in results[:15]: 
            html_body += f"""
                    <div style="border-bottom: 1px solid #1a1a1a; padding-bottom: 15px;">
                        <h2 style="font-size: 18px; margin-bottom: 5px;"><a href="{item['link']}" style="color: #1a1a1a; text-decoration: none; font-weight: bold;">{item['title']}</a></h2>
                        <p style="font-size: 14px; line-height: 1.5;">{item['snippet']}</p>
                        <p style="font-size: 12px; font-weight: bold; text-align: right; color: #555;">— {item.get('source', 'Naver News')} —</p>
                    </div>
            """
            
        html_body += """
                </div>
                <p style="text-align: center; font-size: 12px; margin-top: 30px; border-top: 1px solid #1a1a1a; padding-top: 10px;">
                    &copy; 2026 THE NAVER DAILY ARCHIVE. ALL RIGHTS RESERVED.
                </p>
            </div>
        </div>
        """

        msg = MIMEMultipart()
        msg['From'] = auth_user
        msg['To'] = recipient
        msg['Subject'] = subject
        msg.attach(MIMEText(html_body, 'html'))

        with smtplib.SMTP(server_host, server_port) as server:
            server.starttls()
            server.login(auth_user, auth_pw)
            server.send_message(msg)

        return jsonify({'message': 'Success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    Timer(1, open_browser).start()
    app.run(debug=False, port=5001)
